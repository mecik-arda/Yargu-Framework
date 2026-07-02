import os
import random
import string
import io
from pathlib import Path

SAHTE_RAPOR_SABLONU_TR = """{sirket_adi}
GIZLILIK SEVIYESI: IC KULLANIM

YILLIK FAALIYET RAPORU
{rapor_yili} Donemi

1. YONETICI OZETI
Bu rapor, {sirket_adi} bunyesinde yurutulen {proje_adi} projesinin kapsamli degerlendirmesini icermektedir.
Proje, {baslangic_tarihi} tarihinde baslamis olup {bitis_tarihi} tarihinde basariyla tamamlanmistir.

2. PROJE KAPSAMI
Proje kapsaminda asagidaki faaliyetler gerceklestirilmistir:
- Pazar arastirmasi ve fizibilite calismasi
- Teknik altyapi degerlendirmesi
- Insan kaynaklari planlamasi
- Butce ve finansal projeksiyonlar
- Risk analizi ve azaltma stratejileri

3. FINANSAL OZET
{butce} TL butce ile baslatilan projede, toplam harcama {harcama} TL olarak gerceklesmistir.
Tasarruf orani: %{tasarruf_orani}

4. PERFORMANS DEGERLENDIRMESI
Ekip, {ekip_buyuklugu} kisilik kadro ile {tamamlanma_orani}% tamamlanma oranina ulasmistir.
Proje takvimi dahilinde tum asamalar basariyla gecilmistir.

5. SONUC VE ONERILER
Projenin genel degerlendirmesi olumlu olup, onumuzdeki donemde {gelecek_planlari} hedeflenmektedir.

Raporu Hazirlayan: {hazirlayan}
Onaylayan: {onaylayan}
Tarih: {rapor_tarihi}
"""

SAHTE_RAPOR_SABLONU_EN = """{sirket_adi}
CONFIDENTIALITY: INTERNAL USE ONLY

ANNUAL ACTIVITY REPORT
{rapor_yili} Period

1. EXECUTIVE SUMMARY
This report contains a comprehensive evaluation of the {proje_adi} project under {sirket_adi}.
The project started on {baslangic_tarihi} and was successfully completed on {bitis_tarihi}.

2. PROJECT SCOPE
The following activities were carried out within the project scope:
- Market research and feasibility study
- Technical infrastructure assessment
- Human resources planning
- Budget and financial projections
- Risk analysis and mitigation strategies

3. FINANCIAL SUMMARY
The project was initiated with a budget of {butce} TL, with total expenditure of {harcama} TL.
Savings rate: %{tasarruf_orani}

4. PERFORMANCE EVALUATION
The team of {ekip_buyuklugu} personnel achieved a completion rate of {tamamlanma_orani}%.
All phases were successfully completed within the project timeline.

5. CONCLUSION AND RECOMMENDATIONS
The overall evaluation of the project is positive, with plans for {gelecek_planlari} in the upcoming period.

Prepared by: {hazirlayan}
Approved by: {onaylayan}
Date: {rapor_tarihi}
"""

SIRKET_ADLARI_TR = ["Anadolu Teknoloji A.S.", "Bogazici Yazilim Ltd.", "Ege Danismanlik Grubu", "Marmara Bilisim A.S.", "Turk Telekomunikasyon Entegrasyon"]
SIRKET_ADLARI_EN = ["Anatolian Technologies Inc.", "Bosphorus Software Ltd.", "Aegean Consulting Group", "Marmara IT Solutions", "Turkish Telecommunication Integration"]
PROJE_ADLARI_TR = ["Dijital Donusum", "Bulut Altyapi Modernizasyonu", "Siber Guvenlik Guclendirme", "Veri Analitigi Platformu", "Mobil Entegrasyon"]
PROJE_ADLARI_EN = ["Digital Transformation", "Cloud Infrastructure Modernization", "Cybersecurity Hardening", "Data Analytics Platform", "Mobile Integration"]
AYLAR_TR = ["Ocak", "Subat", "Mart", "Nisan", "Mayis", "Haziran", "Temmuz", "Agustos", "Eylul", "Ekim", "Kasim", "Aralik"]
AYLAR_EN = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
GELECEK_TR = ["yapay zeka entegrasyonu", "Avrupa pazarlarina acilim", "urun gamini genisletme", "Ar-Ge yatirimlarinin artirilmasi", "uluslararasi is birlikleri kurma"]
GELECEK_EN = ["AI integration", "expansion into European markets", "product line diversification", "increased R&D investments", "establishing international partnerships"]


class RAGPoisoner:
    def __init__(self):
        self.son_uretilenler = []

    def sahte_sirket_raporu_olustur(self, dil="tr"):
        if dil == "en":
            sablon = SAHTE_RAPOR_SABLONU_EN
            sirket_adi = random.choice(SIRKET_ADLARI_EN)
            proje_adi = random.choice(PROJE_ADLARI_EN)
            ay = random.choice(AYLAR_EN)
            gelecek = random.choice(GELECEK_EN)
        else:
            sablon = SAHTE_RAPOR_SABLONU_TR
            sirket_adi = random.choice(SIRKET_ADLARI_TR)
            proje_adi = random.choice(PROJE_ADLARI_TR)
            ay = random.choice(AYLAR_TR)
            gelecek = random.choice(GELECEK_TR)
        return sablon.format(
            sirket_adi=sirket_adi,
            rapor_yili=random.randint(2023, 2026),
            proje_adi=proje_adi,
            baslangic_tarihi=f"{random.randint(1, 28)} {ay} {random.randint(2023, 2025)}",
            bitis_tarihi=f"{random.randint(1, 28)} {ay} {random.randint(2025, 2026)}",
            butce=f"{random.randint(1, 50) * 100000:,}",
            harcama=f"{random.randint(1, 50) * 95000:,}",
            tasarruf_orani=random.randint(5, 35),
            ekip_buyuklugu=random.randint(3, 25),
            tamamlanma_orani=random.randint(85, 100),
            gelecek_planlari=gelecek,
            hazirlayan=f"Proje Ekibi",
            onaylayan=f"Yonetim Kurulu",
            rapor_tarihi=f"{random.randint(1, 28)}.{random.randint(1, 12)}.{random.randint(2025, 2026)}"
        )

    def pdf_zehirle(self, gizli_prompt, cikti_yolu, gorunur_icerik=None, dil="tr"):
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.colors import white, black
        except ImportError:
            return self._pdf_zehirle_fpdf(gizli_prompt, cikti_yolu, gorunur_icerik, dil)
        Path(cikti_yolu).parent.mkdir(parents=True, exist_ok=True)
        if gorunur_icerik is None:
            gorunur_icerik = self.sahte_sirket_raporu_olustur(dil)
        c = canvas.Canvas(cikti_yolu, pagesize=A4)
        c.setTitle("Faaliyet Raporu" if dil == "tr" else "Activity Report")
        c.setAuthor("Proje Ekibi")
        c.setSubject(gizli_prompt[:200])
        satirlar = gorunur_icerik.split('\n')
        y = 800
        for satir in satirlar:
            if satir.strip():
                c.setFont("Helvetica", 10)
                c.setFillColor(black)
            else:
                c.setFont("Helvetica", 8)
                c.setFillColor(black)
            c.drawString(50, y, satir[:90])
            y -= 14
            if y < 60:
                c.showPage()
                y = 800
        c.setFont("Helvetica", 8)
        c.setFillColor(white)
        gizli_satirlar = [gizli_prompt[i:i+80] for i in range(0, len(gizli_prompt), 80)]
        gizli_y = 50
        for satir in gizli_satirlar:
            c.drawString(30, gizli_y, satir)
            gizli_y -= 10
        c.save()
        self.son_uretilenler.append({"dosya": cikti_yolu, "tur": "pdf", "teknik": "beyaz_font"})
        return cikti_yolu

    def _pdf_zehirle_fpdf(self, gizli_prompt, cikti_yolu, gorunur_icerik=None, dil="tr"):
        try:
            from fpdf import FPDF
        except ImportError:
            return self._pdf_zehirle_txt_fallback(gizli_prompt, cikti_yolu)
        Path(cikti_yolu).parent.mkdir(parents=True, exist_ok=True)
        if gorunur_icerik is None:
            gorunur_icerik = self.sahte_sirket_raporu_olustur(dil)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=10)
        satirlar = gorunur_icerik.split('\n')
        for satir in satirlar:
            pdf.cell(0, 6, txt=satir[:90], ln=True)
        pdf.set_font("Helvetica", size=6)
        pdf.set_text_color(255, 255, 255)
        pdf.set_y(270)
        gizli_satirlar = [gizli_prompt[i:i+70] for i in range(0, len(gizli_prompt), 70)]
        for satir in gizli_satirlar:
            pdf.cell(0, 4, txt=satir, ln=True)
        pdf.output(cikti_yolu)
        self.son_uretilenler.append({"dosya": cikti_yolu, "tur": "pdf", "teknik": "beyaz_font_fpdf"})
        return cikti_yolu

    def _pdf_zehirle_txt_fallback(self, gizli_prompt, cikti_yolu):
        cikti_yolu = cikti_yolu.replace('.pdf', '_zehirli.txt')
        return self.txt_zehirle(gizli_prompt, cikti_yolu, "zero_width")

    def docx_zehirle(self, gizli_prompt, cikti_yolu, gorunur_icerik=None, dil="tr"):
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
        except ImportError:
            return self._docx_fallback(gizli_prompt, cikti_yolu, gorunur_icerik, dil)
        Path(cikti_yolu).parent.mkdir(parents=True, exist_ok=True)
        if gorunur_icerik is None:
            gorunur_icerik = self.sahte_sirket_raporu_olustur(dil)
        doc = Document()
        baslik = doc.add_heading('Faaliyet Raporu' if dil == 'tr' else 'Activity Report', 0)
        satirlar = gorunur_icerik.split('\n')
        for satir in satirlar:
            if satir.strip():
                doc.add_paragraph(satir[:200])
        gizli_paragraf = doc.add_paragraph()
        gizli_run = gizli_paragraf.add_run(gizli_prompt)
        gizli_run.font.size = Pt(6)
        gizli_run.font.color.rgb = RGBColor(255, 255, 255)
        doc.save(cikti_yolu)
        self.son_uretilenler.append({"dosya": cikti_yolu, "tur": "docx", "teknik": "beyaz_font_gizli"})
        return cikti_yolu

    def _docx_fallback(self, gizli_prompt, cikti_yolu, gorunur_icerik=None, dil="tr"):
        cikti_yolu = cikti_yolu.replace('.docx', '_zehirli.txt')
        return self.txt_zehirle(gizli_prompt, cikti_yolu, "zero_width")

    def txt_zehirle(self, gizli_prompt, cikti_yolu, teknik="zero_width", gorunur_icerik=None, dil="tr"):
        Path(cikti_yolu).parent.mkdir(parents=True, exist_ok=True)
        if gorunur_icerik is None:
            gorunur_icerik = self.sahte_sirket_raporu_olustur(dil)
        if teknik == "zero_width":
            zero_width_chars = ['​', '‌', '‍', '﻿']
            gizli_gizlenmis = ""
            for i, karakter in enumerate(gizli_prompt):
                gizli_gizlenmis += karakter
                if i % 3 == 0:
                    gizli_gizlenmis += random.choice(zero_width_chars)
            icerik = f"{gorunur_icerik}\n\n{gizli_gizlenmis}"
        elif teknik == "rtl_override":
            rtl_override = '‮'
            pop_direction = '‬'
            gizli_ters = gizli_prompt[::-1]
            icerik = f"{gorunur_icerik}\n\n{rtl_override}{gizli_ters}{pop_direction}"
        elif teknik == "satir_sonu":
            uzun_bosluk = " " * 500
            icerik = f"{gorunur_icerik}\n{uzun_bosluk}{gizli_prompt}"
        else:
            icerik = f"{gorunur_icerik}\n\n[GIZLI]{gizli_prompt}[/GIZLI]"
        with open(cikti_yolu, "w", encoding="utf-8") as f:
            f.write(icerik)
        self.son_uretilenler.append({"dosya": cikti_yolu, "tur": "txt", "teknik": teknik})
        return cikti_yolu

    def toplu_uret(self, gizli_prompt, adet=5, formatlar=None, cikti_dizini="cikti/zehirli_dosyalar", dil="tr"):
        if formatlar is None:
            formatlar = ["pdf", "docx", "txt"]
        Path(cikti_dizini).mkdir(parents=True, exist_ok=True)
        uretilenler = []
        for i in range(adet):
            for fmt in formatlar:
                dosya_adi = f"rapor_{i+1}_{random.randint(1000, 9999)}.{fmt}"
                cikti_yolu = os.path.join(cikti_dizini, dosya_adi)
                if fmt == "pdf":
                    self.pdf_zehirle(gizli_prompt, cikti_yolu, dil=dil)
                elif fmt == "docx":
                    self.docx_zehirle(gizli_prompt, cikti_yolu, dil=dil)
                elif fmt == "txt":
                    teknik_secimi = random.choice(["zero_width", "rtl_override", "satir_sonu"])
                    self.txt_zehirle(gizli_prompt, cikti_yolu, teknik=teknik_secimi, dil=dil)
                uretilenler.append(cikti_yolu)
        return uretilenler
